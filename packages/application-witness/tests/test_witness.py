from __future__ import annotations

import hashlib
import json
import stat
import subprocess
import sys
import tempfile
import textwrap
import threading
import time
import unittest
from pathlib import Path
from typing import get_type_hints
from unittest import mock

from docx import Document
from office_application_witness import ApplicationWitness
from office_application_witness.api import _inspect_output, _source_matches
from office_application_witness.contracts import RefusalReason, WitnessRefusal
from openpyxl import Workbook

FAKE_SOFFICE = rf"""#!{sys.executable}
import pathlib, shutil, sys
if "--version" in sys.argv:
    print("LibreOffice Fake 1.0")
    raise SystemExit(0)
out = pathlib.Path(sys.argv[sys.argv.index("--outdir") + 1])
source = pathlib.Path(sys.argv[-1])
out.mkdir(parents=True, exist_ok=True)
if source.suffix.lower() == ".xlsx":
    shutil.copyfile(source, out / (source.stem + ".xlsx"))
else:
    (out / (source.stem + ".pdf")).write_bytes(b"%PDF-1.4\n% fake witness\n")
"""

SLOW_SOFFICE = rf"""#!{sys.executable}
import time
time.sleep(30)
"""

SLOW_SUCCESS_SOFFICE = rf"""#!{sys.executable}
import pathlib, sys, time
if "--version" in sys.argv:
    print("LibreOffice Fake 1.0")
    raise SystemExit(0)
time.sleep(1)
out = pathlib.Path(sys.argv[sys.argv.index("--outdir") + 1])
source = pathlib.Path(sys.argv[-1])
(out / (source.stem + ".pdf")).write_bytes(b"%PDF-1.4\n% fake witness\n")
"""

PROCESS_GROUP_SOFFICE = rf"""#!{sys.executable}
import pathlib, subprocess, sys, time
marker = pathlib.Path(sys.argv[-1]).parent.parent.parent / "child-survived"
subprocess.Popen([sys.executable, "-c", "import pathlib,time;time.sleep(2);pathlib.Path(r'%s').write_text('bad')" % marker])
time.sleep(30)
"""

IGNORE_TERM_GROUP_SOFFICE = rf"""#!{sys.executable}
import pathlib, subprocess, sys, time
marker = pathlib.Path(sys.argv[-1]).parent.parent.parent / "term-ignoring-child-survived"
child = "import pathlib,signal,time;signal.signal(signal.SIGTERM,signal.SIG_IGN);time.sleep(2);pathlib.Path(r'%s').write_text('bad')" % marker
subprocess.Popen([sys.executable, "-c", child])
time.sleep(30)
"""

FIFO_OUTPUT_SOFFICE = rf"""#!{sys.executable}
import os, pathlib, sys
out = pathlib.Path(sys.argv[sys.argv.index("--outdir") + 1])
source = pathlib.Path(sys.argv[-1])
os.mkfifo(out / (source.stem + ".pdf"))
"""

SYMLINK_OUTPUT_SOFFICE = rf"""#!{sys.executable}
import pathlib, sys
out = pathlib.Path(sys.argv[sys.argv.index("--outdir") + 1])
source = pathlib.Path(sys.argv[-1])
out.mkdir(parents=True, exist_ok=True)
(out / (source.stem + ".pdf")).symlink_to(source)
"""

REPLACED_OUTPUT_DIRECTORY_SOFFICE = rf"""#!{sys.executable}
import pathlib, sys
out = pathlib.Path(sys.argv[sys.argv.index("--outdir") + 1])
source = pathlib.Path(sys.argv[-1])
external = source.parents[3] / "escaped-output"
external.mkdir(mode=0o700)
out.rmdir()
out.symlink_to(external, target_is_directory=True)
(out / (source.stem + ".pdf")).write_bytes(b"%PDF-1.4\n% escaped witness\n")
"""

SUCCESS_WITH_LIVE_CHILD_SOFFICE = rf"""#!{sys.executable}
import pathlib, subprocess, sys
out = pathlib.Path(sys.argv[sys.argv.index("--outdir") + 1])
source = pathlib.Path(sys.argv[-1])
marker = source.parents[2] / "child-survived-after-success"
child = "import pathlib,time;time.sleep(1.5);pathlib.Path(r'%s').write_text('bad')" % marker
subprocess.Popen([sys.executable, "-c", child])
(out / (source.stem + ".pdf")).write_bytes(b"%PDF-1.4\n% fake witness\n")
"""

NO_OUTPUT_SOFFICE = rf"""#!{sys.executable}
raise SystemExit(0)
"""


class ApplicationWitnessTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.work = self.root / "work"
        self.fake = self.make_executable("fake-soffice", FAKE_SOFFICE)
        self.witness = ApplicationWitness(self.work, executable=self.fake)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def make_executable(self, name: str, content: str) -> Path:
        path = self.root / name
        path.write_text(textwrap.dedent(content))
        path.chmod(path.stat().st_mode | stat.S_IXUSR)
        return path

    def test_docx_pdf_observation_uses_private_clone_and_cleans_workspace(self) -> None:
        source = self.root / "source.docx"
        document = Document()
        document.add_paragraph("Witness")
        document.save(source)
        before = hashlib.sha256(source.read_bytes()).hexdigest()
        result = self.witness.observe(source, "docx", timeout_seconds=5)
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["schema_version"], 1)
        self.assertEqual(
            set(result),
            {
                "schema_version",
                "status",
                "artifact_type",
                "source_sha256",
                "source_bytes",
                "source_unchanged",
                "witness",
                "private_workspace_artifacts_retained",
                "microsoft_office_equivalence",
                "latency_ms",
            },
        )
        self.assertEqual(result["witness"]["operation"], "pdf_render")
        self.assertEqual(result["witness"]["claim"], "libreoffice_private_clone_observed")
        self.assertTrue(result["source_unchanged"])
        self.assertFalse(result["private_workspace_artifacts_retained"])
        self.assertEqual(result["witness"]["process_isolation"], "trusted_executable_not_sandboxed")
        self.assertEqual(result["microsoft_office_equivalence"], "not_claimed")
        self.assertEqual(hashlib.sha256(source.read_bytes()).hexdigest(), before)
        self.assertEqual(list(self.work.iterdir()), [])

    def test_xlsx_roundtrip_is_observed_but_recalculation_not_overclaimed(self) -> None:
        source = self.root / "source.xlsx"
        workbook = Workbook()
        workbook.active["A1"] = "=1+1"
        workbook.save(source)
        result = self.witness.observe(source, "xlsx", timeout_seconds=5)
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["witness"]["operation"], "recalculation_roundtrip")
        self.assertEqual(result["witness"]["formula_recalculation"], "requested_not_semantically_verified")

    def test_host_supplied_runtime_identity_is_bounded_and_reported(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        identity = {
            "application_version": "LibreOffice 24.2.7.2",
            "image_digest": "sha256:" + "a" * 64,
        }
        result = ApplicationWitness(
            self.work,
            executable=self.fake,
            runtime_identity=identity,
        ).observe(source, "docx", timeout_seconds=5)
        self.assertEqual(result["status"], "ok")
        self.assertEqual(result["witness"]["runtime_identity"], identity)
        self.assertEqual(result["witness"]["version"], "LibreOffice 24.2.7.2")

    def test_runtime_identity_is_closed_and_validated_before_use(self) -> None:
        invalid = (
            {},
            {"application_version": "LibreOffice 24.2", "image_digest": "latest"},
            {"application_version": "x" * 129, "image_digest": "sha256:" + "a" * 64},
            {"application_version": "LibreOffice 24.2", "image_digest": "sha256:" + "a" * 64, "extra": "x"},
        )
        for identity in invalid:
            with self.subTest(identity=identity), self.assertRaises(ValueError):
                ApplicationWitness(self.work, executable=self.fake, runtime_identity=identity)

    def test_runtime_identity_mutation_after_construction_is_refused(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        witness = ApplicationWitness(
            self.work,
            executable=self.fake,
            runtime_identity={
                "application_version": "LibreOffice 24.2.7.2",
                "image_digest": "sha256:" + "a" * 64,
            },
        )
        witness.runtime_identity["image_digest"] = "latest"
        result = witness.observe(source, "docx", timeout_seconds=5)
        self.assertEqual((result["status"], result["reason"]), ("refused", "validation_failure"))

    def test_source_symlink_is_refused(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        link = self.root / "link.docx"
        link.symlink_to(source)
        result = self.witness.observe(link, "docx")
        self.assertEqual(result["status"], "refused")
        self.assertEqual(result["reason"], "validation_failure")

    def test_workdir_must_be_absolute_private_and_symlink_free(self) -> None:
        relative = Path("relative-witness-work")
        with self.assertRaisesRegex(ValueError, "absolute"):
            ApplicationWitness(relative, executable=self.fake)
        self.assertFalse(relative.exists())

        insecure = self.root / "insecure-work"
        insecure.mkdir(mode=0o700)
        insecure.chmod(0o755)
        with self.assertRaisesRegex(ValueError, "private"):
            ApplicationWitness(insecure, executable=self.fake)

        outside = self.root / "outside"
        outside.mkdir(mode=0o700)
        link = self.root / "linked-parent"
        link.symlink_to(outside, target_is_directory=True)
        with self.assertRaisesRegex(ValueError, "invalid witness workdir"):
            ApplicationWitness(link / "child", executable=self.fake)
        self.assertFalse((outside / "child").exists())

    def test_cleanup_failure_cannot_return_success(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        with mock.patch("office_application_witness.api.shutil.rmtree", side_effect=OSError("blocked")):
            result = self.witness.observe(source, "docx", timeout_seconds=5)
        self.assertEqual(result["status"], "refused")
        self.assertEqual(result["reason"], "validation_failure")

    def test_timeout_is_typed_and_private_workspace_is_cleaned(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        slow = self.make_executable("slow-soffice", SLOW_SOFFICE)
        result = ApplicationWitness(self.work, executable=slow).observe(source, "docx", timeout_seconds=1)
        self.assertEqual(result["status"], "refused")
        self.assertEqual(result["reason"], "application_timeout")
        self.assertEqual(result["schema_version"], 1)
        self.assertEqual(list(self.work.iterdir()), [])

    def test_unavailable_application_is_typed(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        result = ApplicationWitness(self.work, executable=self.root / "missing").observe(source, "docx")
        self.assertEqual(result["status"], "refused")
        self.assertEqual(result["reason"], "application_unavailable")

    def test_missing_source_is_validation_failure_not_application_unavailable(self) -> None:
        result = self.witness.observe(self.root / "missing.docx", "docx")
        self.assertEqual((result["status"], result["reason"]), ("refused", "validation_failure"))

    def test_missing_application_output_is_application_failure(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        executable = self.make_executable("no-output-soffice", NO_OUTPUT_SOFFICE)
        result = ApplicationWitness(self.work, executable=executable).observe(source, "docx", timeout_seconds=5)
        self.assertEqual((result["status"], result["reason"]), ("refused", "application_failure"))

    def test_invalid_office_package_is_refused_before_application(self) -> None:
        source = self.root / "bad.docx"
        source.write_bytes(b"not a ZIP package")
        result = self.witness.observe(source, "docx")
        self.assertEqual(result["status"], "refused")
        self.assertEqual(result["reason"], "validation_failure")

    def test_source_change_during_observation_is_stale(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        slow_success = self.make_executable("slow-success-soffice", SLOW_SUCCESS_SOFFICE)
        thread = threading.Thread(
            target=lambda: (time.sleep(0.25), source.write_bytes(source.read_bytes() + b"changed"))
        )
        thread.start()
        result = ApplicationWitness(self.work, executable=slow_success).observe(source, "docx", timeout_seconds=5)
        thread.join()
        self.assertEqual(result["status"], "refused")
        self.assertEqual(result["reason"], "stale_snapshot")

    def test_source_hash_rejects_metadata_change_during_final_read(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        initial = source.stat()
        digest = hashlib.sha256(source.read_bytes()).hexdigest()
        changed = mock.Mock(wraps=initial)
        changed.st_ctime_ns = initial.st_ctime_ns + 1
        with mock.patch("office_application_witness.api.os.fstat", side_effect=[initial, changed]):
            self.assertFalse(_source_matches(source, initial, digest))

    def test_output_hash_rejects_metadata_change_during_read(self) -> None:
        output = self.root / "artifact.pdf"
        output.write_bytes(b"%PDF-1.4\ncontent\n")
        initial = output.stat()
        changed = mock.Mock(wraps=initial)
        changed.st_mtime_ns = initial.st_mtime_ns + 1
        directory = output.parent.stat()
        with mock.patch(
            "office_application_witness.api.os.fstat", side_effect=[directory, initial, changed]
        ), self.assertRaisesRegex(ValueError, "stable"):
            _inspect_output(output, "docx")

    def test_timeout_kills_spawned_process_group(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        group = self.make_executable("group-soffice", PROCESS_GROUP_SOFFICE)
        result = ApplicationWitness(self.work, executable=group).observe(source, "docx", timeout_seconds=1)
        self.assertEqual(result["reason"], "application_timeout")
        time.sleep(2.2)
        self.assertFalse((self.work / "child-survived").exists())

    def test_timeout_kills_child_that_ignores_sigterm_after_leader_exits(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        group = self.make_executable("ignore-term-group-soffice", IGNORE_TERM_GROUP_SOFFICE)
        result = ApplicationWitness(self.work, executable=group).observe(source, "docx", timeout_seconds=1)
        self.assertEqual(result["reason"], "application_timeout")
        time.sleep(2.2)
        self.assertFalse((self.work / "term-ignoring-child-survived").exists())

    def test_success_kills_remaining_process_group_before_cleanup(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        executable = self.make_executable("success-with-live-child-soffice", SUCCESS_WITH_LIVE_CHILD_SOFFICE)
        result = ApplicationWitness(self.work, executable=executable).observe(source, "docx", timeout_seconds=5)
        self.assertEqual(result["status"], "ok")
        time.sleep(1.8)
        self.assertFalse((self.work / "child-survived-after-success").exists())
        self.assertEqual(list(self.work.iterdir()), [])

    def test_fifo_application_output_refuses_without_blocking(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        malicious = self.make_executable("fifo-output-soffice", FIFO_OUTPUT_SOFFICE)
        started = time.monotonic()
        result = ApplicationWitness(self.work, executable=malicious).observe(source, "docx", timeout_seconds=5)
        self.assertLess(time.monotonic() - started, 3)
        self.assertEqual((result["status"], result["reason"]), ("refused", "validation_failure"))
        self.assertEqual(list(self.work.iterdir()), [])

    def test_symlink_application_output_is_refused_and_cleaned(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        malicious = self.make_executable("symlink-soffice", SYMLINK_OUTPUT_SOFFICE)
        result = ApplicationWitness(self.work, executable=malicious).observe(source, "docx", timeout_seconds=5)
        self.assertEqual(result["status"], "refused")
        self.assertEqual(result["reason"], "validation_failure")
        self.assertEqual(list(self.work.iterdir()), [])

    def test_replaced_output_directory_cannot_produce_success(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        malicious = self.make_executable("replaced-output-directory-soffice", REPLACED_OUTPUT_DIRECTORY_SOFFICE)
        result = ApplicationWitness(self.work, executable=malicious).observe(source, "docx", timeout_seconds=5)
        self.assertEqual((result["status"], result["reason"]), ("refused", "validation_failure"))
        self.assertTrue((self.root / "escaped-output" / "artifact.pdf").is_file())
        self.assertEqual(list(self.work.iterdir()), [])

    def test_json_request_cannot_override_host_executable_or_workdir(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        request = {"source": str(source), "artifact_type": "docx", "executable": "/bin/false"}
        completed = subprocess.run(
            [
                sys.executable,
                "-m",
                "office_application_witness",
                "--workdir",
                str(self.work),
                "--executable",
                str(self.fake),
            ],
            input=json.dumps(request),
            text=True,
            capture_output=True,
            check=False,
        )
        result = json.loads(completed.stdout)
        self.assertEqual(result["status"], "refused")
        self.assertEqual(result["reason"], "validation_failure")
        self.assertEqual(completed.returncode, 2)

    def test_cli_uses_host_configuration_and_returns_observation(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        request = {"source": str(source), "artifact_type": "docx", "timeout_seconds": 5}
        completed = subprocess.run(
            [
                sys.executable,
                "-m",
                "office_application_witness",
                "--workdir",
                str(self.work),
                "--executable",
                str(self.fake),
            ],
            input=json.dumps(request),
            text=True,
            capture_output=True,
            check=False,
        )
        self.assertEqual(completed.returncode, 0, completed.stderr)
        self.assertEqual(json.loads(completed.stdout)["status"], "ok")

    def test_json_request_cannot_override_host_runtime_identity(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        request = {
            "source": str(source),
            "artifact_type": "docx",
            "runtime_identity": {
                "application_version": "attacker-selected",
                "image_digest": "sha256:" + "0" * 64,
            },
        }
        completed = subprocess.run(
            [
                sys.executable,
                "-m",
                "office_application_witness",
                "--workdir",
                str(self.work),
                "--executable",
                str(self.fake),
                "--runtime-version",
                "LibreOffice Fake 1.0",
                "--runtime-image-digest",
                "sha256:" + "a" * 64,
            ],
            input=json.dumps(request),
            text=True,
            capture_output=True,
            check=False,
        )
        self.assertEqual(completed.returncode, 2)
        self.assertEqual(json.loads(completed.stdout)["reason"], "validation_failure")

    def test_cli_filesystem_configuration_error_is_typed_without_path_leak(self) -> None:
        source = self.root / "source.docx"
        Document().save(source)
        private_path = "/proc/1/office-witness-private-do-not-log"
        completed = subprocess.run(
            [
                sys.executable,
                "-m",
                "office_application_witness",
                "--workdir",
                private_path,
                "--executable",
                str(self.fake),
            ],
            input=json.dumps({"source": str(source), "artifact_type": "docx"}),
            text=True,
            capture_output=True,
            check=False,
        )
        self.assertEqual(completed.returncode, 2)
        self.assertEqual(json.loads(completed.stdout)["reason"], "validation_failure")
        self.assertNotIn(private_path, completed.stdout + completed.stderr)
        self.assertNotIn("Traceback", completed.stderr)

    def test_refusal_reason_type_is_closed(self) -> None:
        self.assertEqual(get_type_hints(WitnessRefusal)["reason"], RefusalReason)

    def test_direct_api_malformed_source_and_artifact_type_are_typed(self) -> None:
        for source, artifact_type in (([], "docx"), (self.root / "source.docx", [])):
            with self.subTest(source=source, artifact_type=artifact_type):
                result = self.witness.observe(source, artifact_type)  # type: ignore[arg-type]
                self.assertEqual((result["status"], result["reason"]), ("refused", "validation_failure"))


if __name__ == "__main__":
    unittest.main()
